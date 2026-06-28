import base64
import io
from zipfile import ZipFile

from docx import Document

from src.services.ooxml_document_service import apply_node_replacements_to_docx, parse_docx_template


def _docx_bytes_with_activity_table(*, include_image: bool = False) -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("III. HOẠT ĐỘNG DẠY HỌC")
    document.add_paragraph("- Kế hoạch bài dạy, bài giảng Power point.")

    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Hoạt động của giáo viên"
    table.cell(0, 1).text = "Hoạt động của học sinh"
    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(1, 0).text = "1. Khởi động:"
    table.cell(2, 0).text = "- GV nêu nhiệm vụ"
    table.cell(2, 1).text = "- HS thực hiện nhiệm vụ"
    table.cell(3, 0).merge(table.cell(3, 1))
    table.cell(3, 0).text = "2. Hoạt động khám phá:"

    if include_image:
        image_stream = io.BytesIO(base64.b64decode(_ONE_BY_ONE_PNG))
        document.add_picture(image_stream)

    document.save(buffer)
    return buffer.getvalue()


def test_ooxml_parser_marks_structure_fixed_and_body_generated():
    layout, _ = parse_docx_template(_docx_bytes_with_activity_table())
    nodes = list(_iter_nodes(layout))

    heading = next(node for node in nodes if node["text"] == "III. HOẠT ĐỘNG DẠY HỌC")
    bullet = next(node for node in nodes if "Kế hoạch bài dạy" in node["text"])
    table_header = next(node for node in nodes if node["text"] == "Hoạt động của giáo viên")
    phase = next(node for node in nodes if node["text"] == "1. Khởi động:")
    content = next(node for node in nodes if node["text"] == "- GV nêu nhiệm vụ")

    assert heading["role"] == "fixed"
    assert bullet["role"] == "generated"
    assert bullet["display_text"] == ""
    assert table_header["role"] == "fixed"
    assert phase["role"] == "fixed"
    assert content["role"] == "generated"
    assert content["table_context"]["column_header"] == "Hoạt động của giáo viên"
    assert content["table_context"]["phase"] == "1. Khởi động:"


def test_ooxml_parser_does_not_freeze_teacher_body_labels_in_activity_rows():
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("III. HOẠT ĐỘNG DẠY HỌC")
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Hoạt động của giáo viên"
    table.cell(0, 1).text = "Hoạt động của học sinh"
    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(1, 0).text = "1. Khởi động:"
    table.cell(2, 0).text = "Học sinh 1: Đưa ra một con số bất kỳ."
    table.cell(2, 1).text = "- HS trả lời:"
    document.save(buffer)

    layout, _ = parse_docx_template(buffer.getvalue())
    nodes = list(_iter_nodes(layout))
    teacher_body = next(node for node in nodes if node["text"].startswith("Học sinh 1:"))
    student_body = next(node for node in nodes if node["text"].startswith("- HS trả lời"))

    assert teacher_body["role"] == "generated"
    assert teacher_body["display_text"] == ""
    assert student_body["role"] == "generated"
    assert student_body["display_text"] == ""


def test_ooxml_parser_keeps_only_real_structure_fixed_for_body_like_headings():
    buffer = io.BytesIO()
    document = Document()
    paragraphs = [
        "CHƯƠNG I: TẬP HỢP CÁC SỐ TỰ NHIÊN",
        "I.MỤC TIÊU:",
        "2. Năng lực",
        "- Năng lực riêng:",
        "III. TIẾN TRÌNH DẠY HỌC",
        "A. HOẠT ĐỘNG KHỞI ĐỘNG (MỞ ĐẦU)",
        "B.HÌNH THÀNH KIẾN THỨC MỚI",
        "a) Mục tiêu:HS cảm thấy khái niệm tập hợp gần gũi với đời sống hàng ngày.",
        "b) Nội dung: HS quan sát hình ảnh trên màn chiếu.",
        "c) Sản phẩm:",
        "- Bước 1: Chuyển giao nhiệm vụ:",
        "GV cho HS quan sát Hình 1.3 SGK-tr6:",
        "Cách 1: Liệt kê các phần tử của tập hợp:",
        "VD: P = {0; 1; 2; 3 ; 4; 5}",
        "P = {0; 1; 2; 3 ; 4; 5}",
        "A. A = [1; 2; 3; 4]",
        "3 U",
        "+ 18 : XVIII",
    ]
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(buffer)

    layout, _ = parse_docx_template(buffer.getvalue())
    by_text = {node["text"]: node for node in _iter_nodes(layout)}

    for text in paragraphs[:7]:
        assert by_text[text]["role"] == "fixed"
        assert by_text[text]["display_text"] == text

    for text in paragraphs[7:]:
        assert by_text[text]["role"] == "generated"
        assert by_text[text]["display_text"] == ""


def test_ooxml_parser_does_not_treat_numeric_first_table_row_as_header():
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("C. HOẠT ĐỘNG LUYỆN TẬP")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "3 532"
    table.cell(0, 1).text = "3 529"
    table.cell(1, 0).text = "So sánh hai số trên."
    table.cell(1, 1).text = "Điền dấu thích hợp."
    document.save(buffer)

    layout, _ = parse_docx_template(buffer.getvalue())
    nodes = list(_iter_nodes(layout))
    first_value = next(node for node in nodes if node["text"] == "3 532")
    second_value = next(node for node in nodes if node["text"] == "3 529")

    assert first_value["role"] == "generated"
    assert first_value["display_text"] == ""
    assert second_value["role"] == "generated"
    assert second_value["display_text"] == ""


def test_ooxml_apply_preserves_media_files():
    docx_bytes = _docx_bytes_with_activity_table(include_image=True)
    layout, node_map = parse_docx_template(docx_bytes)
    generated_node = next(node for node in _iter_nodes(layout) if node["role"] == "generated")

    output_bytes = apply_node_replacements_to_docx(docx_bytes, node_map, {generated_node["node_id"]: "Nội dung mới"})

    output_document = Document(io.BytesIO(output_bytes))
    assert "Nội dung mới" in "\n".join(paragraph.text for paragraph in output_document.paragraphs)
    with ZipFile(io.BytesIO(output_bytes), "r") as docx_zip:
        assert any(name.startswith("word/media/") for name in docx_zip.namelist())


def _iter_nodes(layout):
    for element in layout["elements"]:
        if element.get("type") == "table":
            for row in element["rows"]:
                for cell in row["cells"]:
                    yield from cell["paragraphs"]
        else:
            yield element


_ONE_BY_ONE_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/lR1q"
    "9QAAAABJRU5ErkJggg=="
)
