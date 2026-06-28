import io

import pytest
from docx import Document

from src.services.auth import create_user


def _minimal_docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("TÊN BÀI DẠY: {{ ten_bai_day }}")
    document.add_paragraph("Môn học: {{ mon_hoc }}")
    document.add_paragraph("- Kế hoạch bài dạy, bài giảng Power point.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Hoạt động của giáo viên"
    table.cell(0, 1).text = "Hoạt động của học sinh"
    table.cell(1, 0).text = "{{ hd_giao_vien }}"
    table.cell(1, 1).text = "{{ hd_hoc_sinh }}"
    document.save(buffer)
    return buffer.getvalue()


def _create_verified_test_user(testing_session_local, email: str, password: str, role: str = "user") -> None:
    with testing_session_local() as db:
        user = create_user(db, email=email, password=password)
        user.role = role
        user.is_email_verified = True
        db.commit()


async def _auth_headers(client, email: str, password: str) -> dict[str, str]:
    response = await client.post("/api/v1/auth/login", json={"email": email, "password": password})
    assert response.status_code == 200
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


@pytest.mark.asyncio
async def test_health(client):
    response = await client.get("/health")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "ok"


@pytest.mark.asyncio
async def test_root_serves_frontend(client):
    response = await client.get("/")
    assert response.status_code == 200
    assert '<div id="root"></div>' in response.text


@pytest.mark.asyncio
async def test_chat_demo_route_removed(client):
    response = await client.post("/api/v1/chat", json={"message": ""})
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_legacy_status_route_removed(client):
    response = await client.get("/api/v1/status")
    assert response.status_code == 404


@pytest.mark.asyncio
async def test_legacy_studio_route_removed(client):
    response = await client.post(
        "/api/v1/studio/compose",
        files={"message": (None, "Soạn giáo án bài Hàm số bậc hai môn Toán lớp 10, 2 tiết")},
    )

    assert response.status_code == 404


@pytest.mark.asyncio
async def test_public_generate_requires_auth(client):
    response = await client.post(
        "/api/v1/lesson/generate",
        files={"prompt": (None, "Soạn giáo án môn Toán lớp 10")},
    )

    assert response.status_code == 401


@pytest.mark.asyncio
async def test_template_preview_requires_auth(client):
    response = await client.post(
        "/api/v1/lesson/template-drafts",
        files={
            "template_file": (
                "template.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 401


@pytest.mark.asyncio
async def test_template_draft_upload_patch_publish_and_list(client, testing_session_local):
    email = "preview-user@example.com"
    password = "password123"
    _create_verified_test_user(testing_session_local, email, password)
    headers = await _auth_headers(client, email, password)

    response = await client.post(
        "/api/v1/lesson/template-drafts",
        headers=headers,
        files={
            "template_file": (
                "template.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 201
    data = response.json()
    assert data["status"] == "draft"
    assert data["summary"]["table_count"] == 1
    assert data["summary"]["node_count"] >= 5
    generated_node = _first_node_with_role(data["layout"], "generated")
    assert generated_node is not None
    assert generated_node["display_text"] == ""

    patch_response = await client.patch(
        f"/api/v1/lesson/template-drafts/{data['id']}/nodes",
        headers=headers,
        json={"nodes": [{"node_id": generated_node["node_id"], "role": "fixed", "text": "Giữ nguyên dòng này"}]},
    )

    assert patch_response.status_code == 200
    patched_node = _node_by_id(patch_response.json()["layout"], generated_node["node_id"])
    assert patched_node["role"] == "fixed"
    assert patched_node["display_text"] == "Giữ nguyên dòng này"

    publish_response = await client.post(
        f"/api/v1/lesson/template-drafts/{data['id']}/publish",
        headers=headers,
        json={"name": "Mẫu test"},
    )

    assert publish_response.status_code == 200
    template = publish_response.json()["template"]
    assert template["name"] == "Mẫu test"
    assert template["visibility"] == "personal"

    list_response = await client.get("/api/v1/lesson/templates", headers=headers)
    assert list_response.status_code == 200
    assert any(item["id"] == template["id"] for item in list_response.json()["templates"])

    detail_response = await client.get(f"/api/v1/lesson/templates/{template['id']}", headers=headers)
    assert detail_response.status_code == 200
    detail = detail_response.json()
    assert detail["id"] == template["id"]
    assert detail["summary"]["node_count"] >= 5


@pytest.mark.asyncio
async def test_public_templates_are_admin_managed_and_user_readable(client, testing_session_local):
    admin_email = "admin@example.com"
    user_email = "teacher@example.com"
    password = "password123"
    _create_verified_test_user(testing_session_local, admin_email, password, role="admin")
    _create_verified_test_user(testing_session_local, user_email, password)
    admin_headers = await _auth_headers(client, admin_email, password)
    user_headers = await _auth_headers(client, user_email, password)

    draft_response = await client.post(
        "/api/v1/lesson/template-drafts",
        headers=admin_headers,
        files={
            "template_file": (
                "public-template.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert draft_response.status_code == 201

    publish_response = await client.post(
        f"/api/v1/lesson/template-drafts/{draft_response.json()['id']}/publish",
        headers=admin_headers,
        json={"name": "Mẫu chung", "visibility": "public"},
    )
    assert publish_response.status_code == 200
    public_template = publish_response.json()["template"]
    assert public_template["visibility"] == "public"
    assert public_template["can_edit"] is True

    user_list_response = await client.get("/api/v1/lesson/templates", headers=user_headers)
    assert user_list_response.status_code == 200
    user_templates = user_list_response.json()["templates"]
    listed_template = next(item for item in user_templates if item["id"] == public_template["id"])
    assert listed_template["visibility"] == "public"
    assert listed_template["can_edit"] is False
    assert listed_template["can_delete"] is False

    user_detail_response = await client.get(f"/api/v1/lesson/templates/{public_template['id']}", headers=user_headers)
    assert user_detail_response.status_code == 200
    assert user_detail_response.json()["id"] == public_template["id"]

    user_patch_response = await client.patch(
        f"/api/v1/lesson/templates/{public_template['id']}",
        headers=user_headers,
        json={"name": "User không được sửa"},
    )
    assert user_patch_response.status_code == 403

    user_delete_response = await client.delete(f"/api/v1/lesson/templates/{public_template['id']}", headers=user_headers)
    assert user_delete_response.status_code == 403

    admin_patch_response = await client.patch(
        f"/api/v1/lesson/templates/{public_template['id']}",
        headers=admin_headers,
        json={"name": "Mẫu chung đã sửa"},
    )
    assert admin_patch_response.status_code == 200
    assert admin_patch_response.json()["name"] == "Mẫu chung đã sửa"

    user_admin_list_response = await client.get("/api/v1/lesson/admin/templates", headers=user_headers)
    assert user_admin_list_response.status_code == 403

    admin_list_response = await client.get("/api/v1/lesson/admin/templates", headers=admin_headers)
    assert admin_list_response.status_code == 200
    assert any(item["id"] == public_template["id"] for item in admin_list_response.json()["templates"])


@pytest.mark.asyncio
async def test_regular_user_cannot_publish_public_template(client, testing_session_local):
    email = "not-admin@example.com"
    password = "password123"
    _create_verified_test_user(testing_session_local, email, password)
    headers = await _auth_headers(client, email, password)

    draft_response = await client.post(
        "/api/v1/lesson/template-drafts",
        headers=headers,
        files={
            "template_file": (
                "template.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert draft_response.status_code == 201

    publish_response = await client.post(
        f"/api/v1/lesson/template-drafts/{draft_response.json()['id']}/publish",
        headers=headers,
        json={"name": "Không được public", "visibility": "public"},
    )
    assert publish_response.status_code == 403


@pytest.mark.asyncio
async def test_public_generate_route_returns_docx_file(client, testing_session_local, monkeypatch):
    async def _fake_generate_template_replacements(*, editable_nodes, **kwargs):
        replacements = {}
        for index, node in enumerate(editable_nodes):
            if index == 0:
                replacements[node["node_id"]] = "TÊN BÀI DẠY: Lực ma sát"
            elif index == 1:
                replacements[node["node_id"]] = "Môn học: Vật lý"
            else:
                replacements[node["node_id"]] = "Nội dung sinh mới"
        return replacements

    monkeypatch.setattr("src.api.routes.lesson.generate_lesson_template_replacements", _fake_generate_template_replacements)
    email = "route-user@example.com"
    password = "password123"
    _create_verified_test_user(testing_session_local, email, password)
    headers = await _auth_headers(client, email, password)

    draft_response = await client.post(
        "/api/v1/lesson/template-drafts",
        headers=headers,
        files={
            "template_file": (
                "template.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert draft_response.status_code == 201
    publish_response = await client.post(
        f"/api/v1/lesson/template-drafts/{draft_response.json()['id']}/publish",
        headers=headers,
        json={"name": "Mẫu generate"},
    )
    assert publish_response.status_code == 200
    template_id = publish_response.json()["template"]["id"]

    response = await client.post(
        "/api/v1/lesson/generate",
        headers=headers,
        files={
            "template_id": (None, str(template_id)),
            "prompt": (None, "Soan giao an bai Luc ma sat mon Vat ly lop 10, 2 tiet."),
            "truong": (None, "THPT Nguyen Trai"),
            "to_chuyen_mon": (None, "To Khoa hoc tu nhien"),
            "giao_vien": (None, "Tran Minh"),
            "lop": (None, "Lop 10"),
            "mon_hoc": (None, "Vat ly"),
            "ten_bai_day": (None, "Luc ma sat"),
            "so_tiet": (None, "2 tiet"),
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    output_document = Document(io.BytesIO(response.content))
    output_text = "\n".join(paragraph.text for paragraph in output_document.paragraphs)
    assert "Lực ma sát" in output_text
    assert "Vật lý" in output_text


def _iter_layout_nodes(layout):
    for element in layout.get("elements", []):
        if element.get("type") == "table":
            for row in element.get("rows", []):
                for cell in row.get("cells", []):
                    yield from cell.get("paragraphs", [])
        elif "node_id" in element:
            yield element


def _first_node_with_role(layout, role):
    return next((node for node in _iter_layout_nodes(layout) if node.get("role") == role), None)


def _node_by_id(layout, node_id):
    return next(node for node in _iter_layout_nodes(layout) if node["node_id"] == node_id)
