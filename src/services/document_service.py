from __future__ import annotations

import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph


def convert_doc_to_docx(path: Path) -> Path:
    """Convert legacy .doc files on Windows; return .docx files unchanged."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return path
    if suffix != ".doc":
        raise RuntimeError("Only .doc and .docx templates are supported")
    if os.name != "nt":
        raise RuntimeError("Legacy .doc conversion is only available on a Windows server with Microsoft Word installed")

    try:
        import win32com.client  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError("win32com is required to convert .doc templates on Windows") from exc

    output_path = path.with_suffix(".docx")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        document = word.Documents.Open(str(path.resolve()))
        document.SaveAs(str(output_path.resolve()), FileFormat=16)
        document.Close()
    finally:
        word.Quit()
    return output_path


def extract_document_elements(path: Path) -> dict[str, str]:
    document = Document(str(path))
    nodes: dict[str, str] = {}

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            nodes[f"p_{index}"] = text

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                for para_index, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text.strip()
                    if text:
                        nodes[f"t_{table_index}_{row_index}_{col_index}_{para_index}"] = text

    return nodes


def extract_document_frame(path: Path) -> dict[str, Any]:
    document = Document(str(path))
    elements: list[dict[str, Any]] = []
    node_count = 0
    editable_count = 0

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        role = _classify_template_text(text)
        editable_count += int(role["editable"])
        node_count += 1
        elements.append(
            {
                "type": "paragraph",
                "node_id": f"p_{index}",
                "text": text,
                **role,
            }
        )

    for table_index, table in enumerate(document.tables):
        table_rows: list[list[dict[str, Any]]] = []
        for row_index, row in enumerate(table.rows):
            rendered_row: list[dict[str, Any]] = []
            for col_index, cell in enumerate(row.cells):
                paragraphs = []
                node_ids = []
                cell_editable = False
                cell_role = "empty"
                for para_index, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    role = _classify_template_text(text, in_table=True, row_index=row_index)
                    paragraphs.append(text)
                    node_ids.append(f"t_{table_index}_{row_index}_{col_index}_{para_index}")
                    cell_editable = cell_editable or role["editable"]
                    if cell_role == "empty" or role["editable"]:
                        cell_role = role["role"]
                    editable_count += int(role["editable"])
                    node_count += 1
                rendered_row.append(
                    {
                        "node_ids": node_ids,
                        "text": "\n".join(paragraphs),
                        "role": cell_role,
                        "editable": cell_editable,
                    }
                )
            table_rows.append(rendered_row)
        elements.append({"type": "table", "table_index": table_index, "rows": table_rows})

    return {
        "summary": {
            "paragraph_count": len([element for element in elements if element["type"] == "paragraph"]),
            "table_count": len(document.tables),
            "node_count": node_count,
            "editable_node_count": editable_count,
        },
        "elements": elements,
    }


def apply_node_replacements(path: Path, replacements: dict[str, Any]) -> None:
    document = Document(str(path))

    for index, paragraph in enumerate(document.paragraphs):
        node_id = f"p_{index}"
        if node_id in replacements:
            _replace_paragraph_text(paragraph, _normalize_replacement(replacements[node_id]))

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                for para_index, paragraph in enumerate(cell.paragraphs):
                    node_id = f"t_{table_index}_{row_index}_{col_index}_{para_index}"
                    if node_id in replacements:
                        _replace_paragraph_text(paragraph, _normalize_replacement(replacements[node_id]))

    document.save(str(path))


def _normalize_replacement(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(item) for item in value)
    return str(value)


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    style_snapshot = _first_run_style(paragraph)
    paragraph.clear()
    parts = text.split("\n")
    for index, part in enumerate(parts):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        _apply_run_style(run, style_snapshot)


def _first_run_style(paragraph: Paragraph) -> dict[str, Any]:
    if not paragraph.runs:
        return {}
    run = paragraph.runs[0]
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "font_name": run.font.name,
        "font_size": run.font.size,
    }


def _apply_run_style(run: Any, style: dict[str, Any]) -> None:
    if not style:
        return
    run.bold = style["bold"]
    run.italic = style["italic"]
    run.underline = style["underline"]
    run.font.name = style["font_name"]
    run.font.size = style["font_size"]


def _classify_template_text(text: str, *, in_table: bool = False, row_index: int | None = None) -> dict[str, Any]:
    compact = " ".join(text.split())
    lowered = compact.lower()
    editable_keywords = (
        "{{",
        "mục tiêu",
        "yêu cầu cần đạt",
        "thiết bị",
        "học liệu",
        "tiến trình",
        "hoạt động",
        "sản phẩm",
        "đánh giá",
        "nhiệm vụ",
        "nội dung",
        "tổ chức thực hiện",
        "dự kiến",
    )
    admin_keywords = ("trường", "tổ", "giáo viên", "môn", "lớp", "bài", "tiết")

    if "{{" in compact:
        return {"role": "placeholder", "editable": True}

    if in_table and row_index == 0 and len(compact) <= 90:
        return {"role": "table_header", "editable": False}

    if ":" in compact and any(keyword in lowered for keyword in admin_keywords):
        return {"role": "field", "editable": True}

    if len(compact) <= 90 and (_looks_like_heading(compact) or compact.endswith(":")):
        return {"role": "heading", "editable": False}

    if any(keyword in lowered for keyword in editable_keywords) or len(compact) > 120:
        return {"role": "content", "editable": True}

    return {"role": "fixed_text", "editable": False}


def _looks_like_heading(text: str) -> bool:
    normalized = text.strip()
    if not normalized:
        return False
    if normalized[:1].isdigit():
        return True
    if normalized.split(".", 1)[0] in {"I", "II", "III", "IV", "V", "A", "B", "C", "D"}:
        return True
    letters = [char for char in normalized if char.isalpha()]
    if letters and sum(char.isupper() for char in letters) / len(letters) > 0.75:
        return True
    return False
